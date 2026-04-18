# templates.py
# ONE JOB: generate platform-specific formatted DOCX resume files

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

def parse_resume_sections(text: str) -> dict:
    # Split raw resume text into sections
    # Returns dict with keys: name, contact, summary, experience, education, skills, projects

    lines = [line.strip() for line in text.split('\n') if line.strip()]

    sections = {
        "name": "",
        "contact": [],
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "other": []
    }

    current_section = "other"
    section_keywords = {
        "experience": ["experience", "work", "employment", "internship"],
        "education": ["education", "qualification", "academic", "degree"],
        "skills": ["skill", "technologies", "technical", "tools", "languages"],
        "projects": ["project", "portfolio", "work sample"],
        "summary": ["summary", "objective", "profile", "about", "overview"]
    }

    for i, line in enumerate(lines):
        line_lower = line.lower()

        # First non-empty line is usually the name
        if i == 0 and len(line.split()) <= 5:
            sections["name"] = line
            continue

        # Check if line is a section header
        found_section = False
        for section, keywords in section_keywords.items():
            if any(kw in line_lower for kw in keywords) and len(line) < 40:
                current_section = section
                found_section = True
                break

        if not found_section:
            # Check if it's contact info
            if "@" in line or any(char.isdigit() for char in line[:15]) or "linkedin" in line_lower or "github" in line_lower:
                if current_section in ["other", "contact"]:
                    sections["contact"].append(line)
            else:
                sections[current_section].append(line)

    return sections


def set_font(run, size=11, bold=False, color=None, font_name="Calibri"):
    # Helper to set font properties on a run
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)  # color is (R, G, B) tuple


def add_section_heading(doc, text, color=(67, 56, 202)):
    # Adds a styled section heading with a bottom border line
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text.upper())
    set_font(run, size=11, bold=True, color=color)

    # Add a separator line after heading
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_before = Pt(0)
    border_para.paragraph_format.space_after = Pt(6)
    border_run = border_para.add_run("─" * 60)
    set_font(border_run, size=8, color=(200, 200, 200))

    return para


def add_bullet(doc, text, indent=True):
    # Adds a bullet point paragraph
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Inches(0.2)
    run = para.add_run(f"• {text}")
    set_font(run, size=10.5)
    return para


def generate_linkedin_docx(sections: dict, optimized_text: str) -> bytes:
    # LinkedIn format: Summary-first, professional, headline focused
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Name — large and bold
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(sections["name"] or "Your Name")
    set_font(name_run, size=20, bold=True, color=(30, 30, 30))

    # LinkedIn headline
    headline_para = doc.add_paragraph()
    headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_run = headline_para.add_run("Software Developer | Open to Opportunities")
    set_font(headline_run, size=11, color=(80, 80, 80))

    # Contact info centered
    if sections["contact"]:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" · ".join(sections["contact"][:4]))
        set_font(contact_run, size=10, color=(100, 100, 100))

    doc.add_paragraph()  # spacer

    # Summary — LinkedIn prioritizes this
    add_section_heading(doc, "Professional Summary", color=(37, 99, 235))
    if sections["summary"]:
        for line in sections["summary"][:5]:
            add_bullet(doc, line)
    else:
        # Extract first 2 lines from optimized text as summary
        lines = [l for l in optimized_text.split('\n') if len(l) > 40]
        summary_text = lines[0] if lines else "Motivated professional seeking opportunities to contribute and grow."
        add_bullet(doc, summary_text)

    # Experience
    add_section_heading(doc, "Experience", color=(37, 99, 235))
    for line in sections["experience"][:10]:
        if line:
            add_bullet(doc, line)

    # Projects
    if sections["projects"]:
        add_section_heading(doc, "Projects", color=(37, 99, 235))
        for line in sections["projects"][:8]:
            if line:
                add_bullet(doc, line)

    # Education
    add_section_heading(doc, "Education", color=(37, 99, 235))
    for line in sections["education"][:5]:
        if line:
            add_bullet(doc, line)

    # Skills
    add_section_heading(doc, "Skills", color=(37, 99, 235))
    skills_text = " · ".join(sections["skills"][:15]) if sections["skills"] else "Python, JavaScript, React, FastAPI"
    skills_para = doc.add_paragraph()
    skills_run = skills_para.add_run(skills_text)
    set_font(skills_run, size=10.5)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_naukri_docx(sections: dict, optimized_text: str) -> bytes:
    # Naukri format: simple single column, experience first, keyword rich
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(sections["name"] or "Your Name")
    set_font(name_run, size=18, bold=True, color=(20, 20, 20))

    # Contact on one line
    if sections["contact"]:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" | ".join(sections["contact"][:4]))
        set_font(contact_run, size=10, color=(80, 80, 80))

    doc.add_paragraph()

    # Objective — Naukri always wants this
    add_section_heading(doc, "Career Objective", color=(220, 38, 38))
    obj_text = sections["summary"][0] if sections["summary"] else "Seeking a challenging position to utilize my technical skills and contribute to organizational growth."
    obj_para = doc.add_paragraph()
    obj_run = obj_para.add_run(obj_text)
    set_font(obj_run, size=10.5)

    # Skills — Naukri ATS scans skills heavily
    add_section_heading(doc, "Technical Skills", color=(220, 38, 38))
    if sections["skills"]:
        for line in sections["skills"][:8]:
            if line:
                add_bullet(doc, line)
    else:
        add_bullet(doc, "Python, Java, JavaScript, React, FastAPI, MongoDB, Docker, Git")

    # Experience
    add_section_heading(doc, "Work Experience", color=(220, 38, 38))
    for line in sections["experience"][:12]:
        if line:
            add_bullet(doc, line)

    # Projects
    if sections["projects"]:
        add_section_heading(doc, "Projects", color=(220, 38, 38))
        for line in sections["projects"][:10]:
            if line:
                add_bullet(doc, line)

    # Education
    add_section_heading(doc, "Education", color=(220, 38, 38))
    for line in sections["education"][:5]:
        if line:
            add_bullet(doc, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_internshala_docx(sections: dict, optimized_text: str) -> bytes:
    # Internshala format: education first, projects prominent, fresher-friendly
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(sections["name"] or "Your Name")
    set_font(name_run, size=18, bold=True, color=(16, 185, 129))

    # Contact centered
    if sections["contact"]:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(sections["contact"][:4]))
        set_font(contact_run, size=10, color=(100, 100, 100))

    doc.add_paragraph()

    # Education FIRST — Internshala is for students/freshers
    add_section_heading(doc, "Education", color=(16, 185, 129))
    for line in sections["education"][:6]:
        if line:
            add_bullet(doc, line)

    # Projects SECOND — most important for freshers
    add_section_heading(doc, "Projects", color=(16, 185, 129))
    if sections["projects"]:
        for line in sections["projects"][:10]:
            if line:
                add_bullet(doc, line)
    else:
        add_bullet(doc, "Add your projects here — this is the most important section for Internshala")

    # Internships / Experience
    add_section_heading(doc, "Internships & Experience", color=(16, 185, 129))
    for line in sections["experience"][:10]:
        if line:
            add_bullet(doc, line)

    # Skills
    add_section_heading(doc, "Skills", color=(16, 185, 129))
    for line in sections["skills"][:8]:
        if line:
            add_bullet(doc, line)

    # Summary / Objective
    if sections["summary"]:
        add_section_heading(doc, "About Me", color=(16, 185, 129))
        for line in sections["summary"][:3]:
            if line:
                add_bullet(doc, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_general_docx(sections: dict, optimized_text: str) -> bytes:
    # General format: clean, ATS-friendly, works for any portal
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(sections["name"] or "Your Name")
    set_font(name_run, size=18, bold=True, color=(30, 30, 30))

    # Contact
    if sections["contact"]:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" | ".join(sections["contact"][:4]))
        set_font(contact_run, size=10.5, color=(80, 80, 80))

    doc.add_paragraph()

    # Summary
    if sections["summary"]:
        add_section_heading(doc, "Professional Summary")
        for line in sections["summary"][:4]:
            if line:
                add_bullet(doc, line)

    # Experience
    add_section_heading(doc, "Experience")
    for line in sections["experience"][:10]:
        if line:
            add_bullet(doc, line)

    # Education
    add_section_heading(doc, "Education")
    for line in sections["education"][:5]:
        if line:
            add_bullet(doc, line)

    # Skills
    add_section_heading(doc, "Skills")
    for line in sections["skills"][:8]:
        if line:
            add_bullet(doc, line)

    # Projects
    if sections["projects"]:
        add_section_heading(doc, "Projects")
        for line in sections["projects"][:8]:
            if line:
                add_bullet(doc, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_resume_docx(
    resume_text: str,
    optimized_text: str,
    platform: str = "general"
) -> bytes:
    # Main function — decides which template to use based on platform

    # Parse resume into sections
    sections = parse_resume_sections(optimized_text or resume_text)

    # Pick correct template
    if platform == "linkedin":
        return generate_linkedin_docx(sections, optimized_text)
    elif platform == "naukri":
        return generate_naukri_docx(sections, optimized_text)
    elif platform == "internshala":
        return generate_internshala_docx(sections, optimized_text)
    else:
        return generate_general_docx(sections, optimized_text)